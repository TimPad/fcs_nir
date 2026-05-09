"""
Парсер входного DOCX → список структурных элементов документа.
Извлекает текст, таблицы, изображения из Word-файла.
Автоматически распознаёт заголовки и специальные разделы.
"""

import re
import io
from typing import List, Optional, Dict, Tuple
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph as DocxParagraph

from elements import (
    Heading, Paragraph, FigureRef, TableElement,
    ListItem, PageBreak, FormulaElement, DocElement,
    auto_number as _auto_number,
)


# ─────────────────────────────────────────────
#  Парсер DOCX
# ─────────────────────────────────────────────
class GostDocxParser:
    """Парсер для извлечения структуры из DOCX файла."""

    # Ключевые слова специальных разделов (без нумерации)
    SPECIAL_SECTIONS = {
        "РЕФЕРАТ", "СОДЕРЖАНИЕ", "ОПРЕДЕЛЕНИЯ",
        "ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ", "ВВЕДЕНИЕ",
        "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "СПИСОК ИСТОЧНИКОВ", "ПРИЛОЖЕНИЯ",
    }

    # Паттерны для распознавания заголовков
    RE_H1_NUMBERED = re.compile(r'^(\d+)\s+(.+)$')
    RE_H2_NUMBERED = re.compile(r'^(\d+\.\d+)\s+(.+)$')
    RE_H3_NUMBERED = re.compile(r'^(\d+\.\d+\.\d+)\s+(.+)$')

    # Паттерны списков
    RE_LIST_BULLET = re.compile(r'^[•\-\*]\s*(.+)$')
    RE_LIST_NUMBERED = re.compile(r'^(\d+)[\.\)]\s*(.+)$')

    def __init__(self):
        self.elements: List[DocElement] = []
        self.images: Dict[str, bytes] = {}
        self.image_counter = 0

    @staticmethod
    def _normalize(s: str) -> str:
        """Нормализация текста: замена спецсимволов, схлопывание пробелов."""
        s = s.replace('\u00a0', ' ').replace('\t', ' ')
        s = re.sub(r' {2,}', ' ', s)
        return s.strip()

    def _is_special_section(self, text: str) -> bool:
        """Проверяет, является ли заголовок специальным разделом."""
        upper = text.upper().strip()
        return upper in self.SPECIAL_SECTIONS or any(
            upper.startswith(s) for s in self.SPECIAL_SECTIONS
        )

    def _detect_heading_level(self, para: DocxParagraph) -> Optional[int]:
        """Определяет уровень заголовка по стилю Word."""
        style_name = para.style.name if para.style else ""
        
        # Проверяем стили заголовков Word
        if style_name in ("Heading 1", "Заголовок 1", "Heading1", "Заголовок1"):
            return 1
        elif style_name in ("Heading 2", "Заголовок 2", "Heading2", "Заголовок2"):
            return 2
        elif style_name in ("Heading 3", "Заголовок 3", "Heading3", "Заголовок3"):
            return 3
        
        # Для специальных разделов (РЕФЕРАТ, СОДЕРЖАНИЕ, ВВЕДЕНИЕ и т.д.) - всегда уровень 1
        text = self._normalize(para.text).upper()
        if text.strip() in self.SPECIAL_SECTIONS or any(
            text.strip().startswith(s) for s in self.SPECIAL_SECTIONS
        ):
            return 1
        
        # Альтернативно: проверяем форматирование (жирный, центрирование, размер)
        if para.runs:
            first_run = para.runs[0]
            is_bold = first_run.bold
            is_centered = para.alignment is not None and str(para.alignment) == 'WD_ALIGN_PARAGRAPH.CENTER'
            font_size = first_run.font.size.pt if first_run.font.size else 14
            
            # Эвристика: жирный + центрированный + крупный шрифт = заголовок 1
            if is_bold and is_centered and font_size >= 14:
                return 1
            # Жирный + слева = заголовок 2
            elif is_bold and not is_centered and font_size >= 14:
                return 2
        
        return None

    def _extract_image(self, shape, doc_part) -> Optional[Tuple[str, bytes]]:
        """Извлекает изображение из inline shape."""
        try:
            if hasattr(shape, 'image') and shape.image:
                self.image_counter += 1
                filename = f"image_{self.image_counter:03d}.png"
                image_data = shape.image.blob
                return filename, image_data
        except Exception:
            pass
        return None

    def _parse_table(self, table: Table) -> TableElement:
        """Парсит таблицу из DOCX."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = ""
                for para in cell.paragraphs:
                    cell_text += self._normalize(para.text) + " "
                cells.append(self._normalize(cell_text))
            if cells and any(c for c in cells):  # Пропускаем пустые строки
                rows.append(cells)
        
        return TableElement(rows=rows, caption="Таблица", has_header=True)

    def parse(self, docx_bytes: bytes) -> Tuple[List[DocElement], Dict[str, bytes]]:
        """
        Парсит DOCX файл.
        Возвращает кортеж: (список элементов, словарь изображений).
        """
        self.elements = []
        self.images = {}
        self.image_counter = 0

        doc = Document(io.BytesIO(docx_bytes))
        
        # Извлекаем встроенные изображения из документа через отношения
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    target = rel.target_part
                    if hasattr(target, 'blob'):
                        self.image_counter += 1
                        ext = target.filename.split('.')[-1] if target.filename else 'png'
                        filename = f"image_{self.image_counter:03d}.{ext}"
                        self.images[filename] = target.blob
                except Exception:
                    continue

        # Парсим параграфы
        i = 0
        paragraphs = list(doc.paragraphs)
        
        while i < len(paragraphs):
            para = paragraphs[i]
            text = self._normalize(para.text)
            
            # Пропускаем пустые параграфы
            if not text:
                i += 1
                continue

            # Определяем тип элемента
            
            # 1. Заголовок по стилю
            heading_level = self._detect_heading_level(para)
            if heading_level:
                # Пытаемся извлечь номер из текста
                m1 = self.RE_H1_NUMBERED.match(text)
                m2 = self.RE_H2_NUMBERED.match(text)
                m3 = self.RE_H3_NUMBERED.match(text)
                
                if heading_level == 1 and m1:
                    elements_num = m1.group(1)
                    title_text = m1.group(2)
                elif heading_level == 2 and m2:
                    elements_num = m2.group(1)
                    title_text = m2.group(2)
                elif heading_level == 3 and m3:
                    elements_num = m3.group(1)
                    title_text = m3.group(2)
                else:
                    elements_num = ""
                    title_text = text
                
                self.elements.append(Heading(
                    level=heading_level,
                    text=title_text,
                    number=elements_num
                ))
                i += 1
                continue

            # 2. Заголовок по паттерну (если стиль не задан, но текст похож на заголовок)
            # Проверяем: начинается с цифры, весь текст в верхнем регистре, короткая строка
            if len(text) < 150 and (text.isupper() or self.RE_H1_NUMBERED.match(text)):
                m = self.RE_H1_NUMBERED.match(text)
                if m:
                    # Это может быть заголовок раздела
                    potential_title = m.group(2)
                    if self._is_special_section(potential_title) or potential_title.isupper():
                        self.elements.append(Heading(
                            level=1,
                            text=potential_title,
                            number=m.group(1)
                        ))
                        i += 1
                        continue

            # 3. Списки
            m_bullet = self.RE_LIST_BULLET.match(text)
            m_numbered = self.RE_LIST_NUMBERED.match(text)
            
            if m_bullet:
                self.elements.append(ListItem(
                    text=m_bullet.group(1),
                    ordered=False
                ))
                i += 1
                continue
            
            if m_numbered:
                self.elements.append(ListItem(
                    text=m_numbered.group(2),
                    ordered=True,
                    number=int(m_numbered.group(1))
                ))
                i += 1
                continue

            # 4. Обычный абзац
            self.elements.append(Paragraph(text=text))
            i += 1

        # Обрабатываем таблицы отдельно (добавляем после соответствующих абзацев)
        for table in doc.tables:
            table_elem = self._parse_table(table)
            # Вставляем таблицу после последнего абзаца перед ней
            # Для простоты добавляем в конец (можно улучшить логику позиционирования)
            self.elements.append(table_elem)

        # Добавляем плейсхолдеры для изображений
        # Изображения уже извлечены в self.images, теперь создаём ссылки на них
        for idx, filename in enumerate(self.images.keys(), 1):
            # Находим подходящее место для вставки (после первого абзаца или заголовка)
            # Для простоты добавляем после первого Paragraph
            insert_pos = 1
            for j, el in enumerate(self.elements):
                if isinstance(el, Paragraph):
                    insert_pos = j + 1
                    break
            
            self.elements.insert(insert_pos, FigureRef(
                path=filename,
                caption=f"Изображение {idx}",
                image_data=self.images.get(filename)
            ))

        return self.elements, self.images

    def auto_number(self, elements: List) -> List:
        """Сквозная нумерация (делегируется в elements.auto_number)."""
        return _auto_number(elements)
