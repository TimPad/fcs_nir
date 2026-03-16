"""
Генератор DOCX строго по ГОСТ 7.32-2017.
"""

import io
import os
import sys
from typing import List, Optional, Dict

# Streamlit Cloud: добавляем папку скрипта в путь поиска модулей
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from gost_styles import (
    PAGE, BODY, HEADING1, HEADING2, HEADING3,
    FIGURE_CAPTION, TABLE_CAPTION, TABLE_CELL, FOOTER, LIST_ITEM
)
from parser import (
    Heading, Paragraph, FigureRef, TableElement,
    ListItem, PageBreak, SpecialSection, FormulaElement
)


class GostDocxBuilder:

    def __init__(self, meta: Dict):
        self.meta = meta  # title, org, authors, year, theme_code, udc
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    # ──────────────────────────────────────────
    #  СТРАНИЦА
    # ──────────────────────────────────────────
    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width  = PAGE["width"]
        section.page_height = PAGE["height"]
        section.left_margin   = PAGE["margin_left"]
        section.right_margin  = PAGE["margin_right"]
        section.top_margin    = PAGE["margin_top"]
        section.bottom_margin = PAGE["margin_bottom"]

    # ──────────────────────────────────────────
    #  СТИЛИ
    # ──────────────────────────────────────────
    def _setup_styles(self):
        styles = self.doc.styles

        # Стиль Normal
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.first_line_indent = Cm(0)
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after  = Pt(0)

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.first_line_indent = Cm(1.25)
        h2.paragraph_format.space_before = Pt(0)
        h2.paragraph_format.space_after  = Pt(0)

        # Heading 3
        h3 = styles["Heading 3"]
        h3.font.name = "Times New Roman"
        h3.font.size = Pt(14)
        h3.font.bold = True
        h3.font.italic = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.first_line_indent = Cm(1.25)
        h3.paragraph_format.space_before = Pt(0)
        h3.paragraph_format.space_after  = Pt(0)

    # ──────────────────────────────────────────
    #  ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
    # ──────────────────────────────────────────
    def _apply_body_fmt(self, para):
        """Применяет форматирование основного текста."""
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25)
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

    def _set_run_font(self, run, size=14, bold=False, italic=False, name="Times New Roman"):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        # Для кириллицы важно задать eastAsia
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def _add_line_spacing(self, para):
        fmt = para.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)

    # ──────────────────────────────────────────
    #  КОЛОНТИТУЛ
    # ──────────────────────────────────────────
    def _add_footer(self, section=None):
        """Нижний колонтитул: номер страницы по центру."""
        sec = section or self.doc.sections[0]
        footer = sec.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)

        # Поле номера страницы
        run = para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    # ──────────────────────────────────────────
    #  ТИТУЛЬНЫЙ ЛИСТ
    # ──────────────────────────────────────────
    def add_title_page(self):
        m = self.meta
        doc = self.doc

        def centered(text, size=14, bold=False, space_after=0):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(text)
            self._set_run_font(run, size=size, bold=bold)
            return p

        # Министерство / организация
        centered(m.get("ministry", ""), 14, False)
        centered(m.get("org", "Организация"), 14, True)
        centered("", 14, False, space_after=12)

        # УДК / код темы
        p_udc = doc.add_paragraph()
        p_udc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_udc.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(p_udc)
        r = p_udc.add_run(f"УДК {m.get('udc', '')}   Инв. №{m.get('inv_number', '')}")
        self._set_run_font(r, 14)

        centered("", 14)
        centered("УТВЕРЖДАЮ", 14, True)
        p_pos = doc.add_paragraph()
        p_pos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_pos.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(p_pos)
        r = p_pos.add_run(m.get("approver_position", "Руководитель"))
        self._set_run_font(r, 14)
        p_sign = doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sign.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(p_sign)
        r = p_sign.add_run(f"____________ {m.get('approver_name', '')}")
        self._set_run_font(r, 14)
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(p_date)
        r = p_date.add_run(f"«____» ___________ {m.get('year', '20__')} г.")
        self._set_run_font(r, 14)

        centered("", 14, space_after=12)

        # Отчёт о НИР
        centered("ОТЧЁТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ", 14, True)
        centered("", 14)
        centered(m.get("title", "Название НИР"), 14, True)
        centered("", 14)
        if m.get("theme_code"):
            centered(f"(Шифр: {m['theme_code']})", 14, False)

        centered("", 14, space_after=24)

        # Руководитель
        def two_col(left, right, size=14):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._add_line_spacing(p)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(8))
            r = p.add_run(f"{left}\t{right}")
            self._set_run_font(r, size)

        two_col("Руководитель НИР,", "")
        two_col(m.get("head_position", ""), "____________")
        two_col("", m.get("head_name", ""))

        doc.add_paragraph()

        # Исполнители
        authors = m.get("authors", [])
        if authors:
            p_auth = doc.add_paragraph()
            p_auth.paragraph_format.first_line_indent = Cm(0)
            self._add_line_spacing(p_auth)
            r = p_auth.add_run("Исполнители:")
            self._set_run_font(r, 14, bold=True)
            for author in authors:
                two_col(author.get("position", ""), author.get("name", ""))

        centered("", 14, space_after=24)
        centered(f"{m.get('city', 'Город')} {m.get('year', '20__')}", 14, False)

        # Разрыв страницы после титульного листа
        doc.add_page_break()

    # ──────────────────────────────────────────
    #  РЕФЕРАТ (обязательный структурный элемент)
    # ──────────────────────────────────────────
    def add_abstract_placeholder(self):
        doc = self.doc
        p = doc.add_paragraph()
        p.style = "Heading 1"
        p.paragraph_format.page_break_before = True
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("РЕФЕРАТ")
        self._set_run_font(run, 14, bold=True)

        body = doc.add_paragraph()
        self._apply_body_fmt(body)
        r = body.add_run(
            "Отчёт __ с., __ рис., __ табл., __ источников, __ приложений.\n\n"
            "КЛЮЧЕВЫЕ СЛОВА: ___, ___, ___\n\n"
            "Текст реферата."
        )
        self._set_run_font(r, 14)
        doc.add_page_break()

    # ──────────────────────────────────────────
    #  СОДЕРЖАНИЕ (TOC)
    # ──────────────────────────────────────────
    def add_toc(self):
        doc = self.doc

        p = doc.add_paragraph()
        p.style = "Heading 1"
        p.paragraph_format.page_break_before = False
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("СОДЕРЖАНИЕ")
        self._set_run_font(run, 14, bold=True)

        # Поле TOC
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(para)
        run = para.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        instrText.set(qn("xml:space"), "preserve")
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        doc.add_page_break()

    # ──────────────────────────────────────────
    #  ЭЛЕМЕНТЫ ТЕЛА ДОКУМЕНТА
    # ──────────────────────────────────────────
    def add_heading(self, el: Heading):
        doc = self.doc
        level_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        p = doc.add_paragraph(style=level_map.get(el.level, "Heading 1"))

        # Разрыв перед разделами 1-го уровня
        if el.level == 1:
            p.paragraph_format.page_break_before = True
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p.paragraph_format.page_break_before = False
            p.paragraph_format.first_line_indent = Cm(1.25)

        self._add_line_spacing(p)

        text = f"{el.number} {el.text}" if el.number else el.text
        run = p.add_run(text.upper() if el.level == 1 else text)
        bold = (el.level in (1, 2))
        italic = (el.level == 3)
        self._set_run_font(run, 14, bold=bold, italic=italic)

    def add_paragraph(self, el: Paragraph):
        p = self.doc.add_paragraph()
        self._apply_body_fmt(p)
        run = p.add_run(el.text)
        self._set_run_font(run, 14)

    def add_figure(self, el: FigureRef, images: Dict[str, bytes]):
        doc = self.doc
        num = getattr(el, "_number", "?")
        caption_text = f"Рисунок {num}" + (f" — {el.caption}" if el.caption else "")

        # Изображение
        img_data = images.get(el.path) or images.get(os.path.basename(el.path))
        if img_data:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.first_line_indent = Cm(0)
            run = p_img.add_run()
            img_stream = io.BytesIO(img_data)
            run.add_picture(img_stream, width=Cm(14))
        else:
            # Заглушка — рамка с текстом
            p_ph = doc.add_paragraph()
            p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ph.paragraph_format.first_line_indent = Cm(0)
            self._add_line_spacing(p_ph)
            r = p_ph.add_run(f"[ Рисунок: {el.path} ]")
            self._set_run_font(r, 12, italic=True)

        # Подпись под рисунком
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.first_line_indent = Cm(0)
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after  = Pt(6)
        p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p_cap.add_run(caption_text)
        self._set_run_font(run, 14)

    def add_table(self, el: TableElement):
        doc = self.doc
        num = getattr(el, "_number", "?")

        # Подпись над таблицей
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_cap.paragraph_format.first_line_indent = Cm(0)
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after  = Pt(3)
        p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        caption_text = f"Таблица {num} — {el.caption}" if el.caption else f"Таблица {num}"
        run = p_cap.add_run(caption_text)
        self._set_run_font(run, 14)

        if not el.rows:
            return

        ncols = max(len(r) for r in el.rows)
        table = doc.add_table(rows=len(el.rows), cols=ncols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(el.rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= ncols:
                    break
                cell = row.cells[c_idx]
                cell.text = cell_text
                for para in cell.paragraphs:
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)
                    for run in para.runs:
                        self._set_run_font(run, 12,
                                           bold=(r_idx == 0 and el.has_header))

    def add_list_item(self, el: ListItem, counter: int):
        doc = self.doc
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.left_indent   = Cm(1.25)
        fmt.first_line_indent = Cm(-0.5)
        fmt.space_before  = Pt(0)
        fmt.space_after   = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        prefix = f"{counter})" if el.ordered else "–"
        run = p.add_run(f"{prefix} {el.text}")
        self._set_run_font(run, 14)

    def add_formula(self, el: FormulaElement):
        doc = self.doc
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        self._add_line_spacing(p)
        text = el.text
        if el.number:
            text += f"\t({el.number})"
        run = p.add_run(text)
        self._set_run_font(run, 14, italic=True)

    def add_page_break(self):
        self.doc.add_page_break()

    # ──────────────────────────────────────────
    #  ФИНАЛИЗАЦИЯ
    # ──────────────────────────────────────────
    def finalize(self):
        """Добавляем колонтитулы ко всем секциям."""
        for section in self.doc.sections:
            self._add_footer(section)
            section.different_first_page_header_footer = True  # первая стр. без номера

    def build(self, elements: List, images: Dict[str, bytes]) -> bytes:
        """Сборка всего документа, возвращает bytes."""
        list_counter = 0
        prev_was_list = False

        for el in elements:
            if isinstance(el, Heading):
                list_counter = 0
                prev_was_list = False
                self.add_heading(el)

            elif isinstance(el, Paragraph):
                list_counter = 0
                prev_was_list = False
                self.add_paragraph(el)

            elif isinstance(el, FigureRef):
                list_counter = 0
                prev_was_list = False
                self.add_figure(el, images)

            elif isinstance(el, TableElement):
                list_counter = 0
                prev_was_list = False
                self.add_table(el)

            elif isinstance(el, ListItem):
                if not prev_was_list or not el.ordered:
                    list_counter = 0
                list_counter += 1
                self.add_list_item(el, list_counter)
                prev_was_list = True

            elif isinstance(el, FormulaElement):
                list_counter = 0
                prev_was_list = False
                self.add_formula(el)

            elif isinstance(el, PageBreak):
                list_counter = 0
                prev_was_list = False
                self.add_page_break()

        self.finalize()

        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()
